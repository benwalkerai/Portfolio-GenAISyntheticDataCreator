import io
import tempfile
from docx import Document
from docx.shared import Inches
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False
    print('  Warning: ReportLab not installed. PDF generation will be disabled.')

class DocumentGenerator:

    def __init__(self, data_generator):
        self.data_generator = data_generator

    def clear_cache(self):
        if hasattr(self, 'llm_cache'):
            self.llm_cache.clear()
        print('DocumentGenerator cache cleared')

    def generate_document_content(self, content_type, pages, subject):
        target_words = pages * 275
        prompts = {'whitepaper': f"Write a technical whitepaper on {subject} ({pages} pages, ~{target_words} words).\n\nSTRUCTURE:\n- Executive Summary\n- Introduction and Background\n- Technical Analysis and Methodology\n- Findings and Results\n- Conclusions\n\nWrite {pages} pages of professional content with technical depth. Add disclaimer: AI-generated demo content.", 'article': f"Write an article about {subject} ({pages} pages, ~{target_words} words).\n\nSTRUCTURE:\n- Introduction\n- Main content with examples and case studies\n- Conclusion\n\nWrite {pages} pages with practical examples. Add disclaimer: AI-generated demo content.", 'report': f"Write a business report on {subject} ({pages} pages, ~{target_words} words).\n\nSTRUCTURE:\n- Executive Summary\n- Introduction and Methodology\n- Analysis and Data\n- Recommendations\n- Conclusions\n\nWrite {pages} pages with data analysis and strategic insights. Add disclaimer: AI-generated demo content.", 'proposal': f"Write a project proposal for {subject} ({pages} pages, ~{target_words} words).\n\nSTRUCTURE:\n- Project Overview\n- Scope and Requirements\n- Timeline and Budget\n- Implementation Plan\n- Expected Outcomes\n\nWrite {pages} pages covering scope, timeline, and budget. Add disclaimer: AI-generated demo content.", 'design': f"Write a design document for {subject} ({pages} pages, ~{target_words} words).\n\nSTRUCTURE:\n- System Overview\n- Technical Requirements\n- Architecture and Data Flow\n- Security and Performance\n- Deployment\n\nWrite {pages} pages with technical specifications. Add disclaimer: AI-generated demo content."}
        if content_type not in prompts:
            raise ValueError(f'Invalid content type: {content_type}. Choose from {list(prompts.keys())}.')
        prompt = prompts.get(content_type)
        system_prompt = 'You are a technical writer. Write clear, well-structured content in Markdown format. Do NOT output JSON or code blocks.'
        return self.data_generator.generate_with_ollama(prompt, max_tokens=6000, system_prompt=system_prompt)

    def generate_document_content_iterative(self, content_type, pages, subject):
        print(f'Generating {pages}-page {content_type}...')
        sections = []
        section_configs = {'whitepaper': [('Executive Summary', 'executive summary'), ('Introduction', 'introduction and background'), ('Technical Analysis', 'technical analysis'), ('Methodology', 'methodology'), ('Implementation', 'implementation details'), ('Results', 'results and findings'), ('Future Work', 'future implications'), ('Conclusions', 'conclusions and references')], 'article': [('Introduction', 'introduction'), ('Analysis', 'main analysis'), ('Case Studies', 'case studies and examples'), ('Industry Impact', 'industry impact'), ('Current Trends', 'current trends'), ('Best Practices', 'best practices'), ('Future Outlook', 'future outlook'), ('Conclusion', 'conclusion')], 'report': [('Executive Summary', 'executive summary'), ('Introduction', 'introduction and methodology'), ('Market Analysis', 'market analysis'), ('Data Analysis', 'data analysis'), ('Recommendations', 'strategic recommendations'), ('Risk Assessment', 'risk assessment'), ('Implementation', 'implementation plan'), ('Financial Analysis', 'financial analysis'), ('Next Steps', 'conclusions and next steps')], 'proposal': [('Overview', 'project overview'), ('Scope', 'scope and requirements'), ('Timeline', 'timeline and milestones'), ('Budget', 'budget and resources'), ('Implementation', 'implementation strategy'), ('Risk Analysis', 'risk analysis'), ('Team', 'team and expertise'), ('Quality', 'quality assurance'), ('Outcomes', 'expected outcomes')], 'design': [('Overview', 'system overview'), ('Requirements', 'technical requirements'), ('Interface Design', 'interface design'), ('Data Architecture', 'data architecture'), ('Security', 'security and performance'), ('Implementation', 'implementation details'), ('Testing', 'testing and QA'), ('Deployment', 'deployment'), ('Maintenance', 'maintenance')]}
        all_sections = section_configs.get(content_type, section_configs['article'])
        sections_to_use = min(len(all_sections), max(pages, 3))
        selected_sections = all_sections[:sections_to_use]
        while len(selected_sections) < pages:
            additional_sections = [('Additional Analysis', 'additional analysis'), ('Supplementary Research', 'supplementary research'), ('Extended Examples', 'extended examples'), ('Advanced Topics', 'advanced topics')]
            selected_sections.extend(additional_sections[:pages - len(selected_sections)])
        total_target_words = pages * 300
        words_per_section = total_target_words // len(selected_sections)
        for i, (section_title, section_instruction) in enumerate(selected_sections):
            section_prompt = f'Write {section_instruction} for a {content_type} about {subject}. Target {words_per_section} words. Section {i + 1} of {len(selected_sections)}.'
            print(f' Generating section {i + 1}/{len(selected_sections)}: {section_title}')
            system_prompt = 'You are a technical writer. Write clear content in Markdown format. No JSON or code blocks.'
            section_content = self.data_generator.generate_with_ollama(section_prompt, max_tokens=2000, system_prompt=system_prompt)
            formatted_section = f'# {section_title}\n\n{section_content}'
            sections.append(formatted_section)
        final_note = '\n\n# Disclaimer\n\nAI-generated demo content. Not for production use.'
        sections.append(final_note)
        full_content = '\n\n'.join(sections)
        word_count = len(full_content.split())
        print(f' Generated document with {len(sections) - 1} sections, approximately {word_count} words')
        return full_content

    def _clean_content(self, content):
        content = content.replace('```markdown', '').replace('```json', '').replace('```', '')
        lines = content.split('\n')
        if len(lines) > 0 and (lines[0].lower().startswith('here is') or lines[0].lower().startswith('sure, here')):
            content = '\n'.join(lines[1:])
        content = content.strip()
        if content.startswith('{') and '"content":' in content:
            try:
                import json
                data = json.loads(content)
                if 'content' in data:
                    return data['content']
            except:
                pass
        return content

    def _strip_markdown(self, text):
        import re
        text = re.sub('\\*\\*(.*?)\\*\\*', '\\1', text)
        text = re.sub('\\*(.*?)\\*', '\\1', text)
        text = re.sub('`(.*?)`', '\\1', text)
        return text

    def _markdown_to_html(self, text):
        import re
        text = re.sub('\\*\\*(.*?)\\*\\*', '<b>\\1</b>', text)
        text = re.sub('\\*(.*?)\\*', '<i>\\1</i>', text)
        text = re.sub('`(.*?)`', '<font name="Courier">\\1</font>', text)
        return text

    def create_word_document(self, content):
        content = self._clean_content(content)
        doc = Document()
        lines = content.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            if not line:
                i += 1
                continue
            if line.startswith('# '):
                heading_text = line.replace('# ', '', 1)
                heading_text = self._strip_markdown(heading_text)
                doc.add_heading(heading_text, 1)
            elif line.startswith('## '):
                heading_text = line.replace('## ', '', 1)
                heading_text = self._strip_markdown(heading_text)
                doc.add_heading(heading_text, 2)
            elif line.startswith('### '):
                heading_text = line.replace('### ', '', 1)
                heading_text = self._strip_markdown(heading_text)
                doc.add_heading(heading_text, 3)
            elif line.startswith('- ') or line.startswith('* '):
                list_text = line[2:].strip()
                p = doc.add_paragraph(style='List Bullet')
                self._add_formatted_text(p, list_text)
            elif line and line[0].isdigit() and ('. ' in line[:4]):
                list_text = line.split('. ', 1)[1] if '. ' in line else line
                p = doc.add_paragraph(style='List Number')
                self._add_formatted_text(p, list_text)
            else:
                p = doc.add_paragraph()
                self._add_formatted_text(p, line)
            i += 1
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        return doc_bytes.getvalue()

    def _add_formatted_text(self, paragraph, text):
        import re
        parts = re.split('(\\*\\*.*?\\*\\*|\\*.*?\\*|`.*?`)', text)
        for part in parts:
            if not part:
                continue
            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*') and (not part.startswith('**')):
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            elif part.startswith('`') and part.endswith('`'):
                run = paragraph.add_run(part[1:-1])
                run.font.name = 'Courier New'
            else:
                paragraph.add_run(part)

    def create_pdf_document(self, content):
        if not REPORTLAB_AVAILABLE:
            raise ImportError('ReportLab library is not installed. Please install it with: pip install reportlab')
        content = self._clean_content(content)
        pdf_bytes = io.BytesIO()
        doc = SimpleDocTemplate(pdf_bytes, pagesize=letter)
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'], fontSize=18, spaceAfter=30)
        heading_style = ParagraphStyle('CustomHeading', parent=styles['Heading2'], fontSize=14, spaceAfter=12)
        subheading_style = ParagraphStyle('CustomSubheading', parent=styles['Heading3'], fontSize=12, spaceAfter=8)
        story = []
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                story.append(Spacer(1, 12))
                continue
            if line.startswith('# '):
                heading_text = line.replace('# ', '', 1)
                heading_text = self._strip_markdown(heading_text)
                story.append(Paragraph(heading_text, heading_style))
                story.append(Spacer(1, 6))
            elif line.startswith('## '):
                heading_text = line.replace('## ', '', 1)
                heading_text = self._strip_markdown(heading_text)
                story.append(Paragraph(heading_text, subheading_style))
                story.append(Spacer(1, 4))
            else:
                formatted_text = self._markdown_to_html(line)
                story.append(Paragraph(formatted_text, styles['Normal']))
                story.append(Spacer(1, 12))
        doc.build(story)
        pdf_bytes.seek(0)
        return pdf_bytes.getvalue()

    def _strip_markdown(self, text):
        import re
        text = re.sub('\\*\\*(.*?)\\*\\*', '\\1', text)
        text = re.sub('\\*(.*?)\\*', '\\1', text)
        text = re.sub('`(.*?)`', '\\1', text)
        return text

    def _markdown_to_html(self, text):
        import re
        text = re.sub('\\*\\*(.*?)\\*\\*', '<b>\\1</b>', text)
        text = re.sub('\\*(.*?)\\*', '<i>\\1</i>', text)
        text = re.sub('`(.*?)`', '<font name="Courier">\\1</font>', text)
        return text

    def create_text_document(self, content):
        content = self._clean_content(content)
        text_content = content.replace('# ', '')
        text_content = text_content.replace('## ', '')
        header = '=' * 80 + '\n'
        header += 'SYNTHETIC DOCUMENT - GENERATED CONTENT\n'
        header += '=' * 80 + '\n\n'
        final_content = header + text_content
        return final_content.encode('utf-8')

    def generate_document(self, content_type, pages, subject, file_format):
        if pages >= 3:
            content = self.generate_document_content_iterative(content_type, pages, subject)
            print(f' Used iterative approach for {pages} pages')
        else:
            content = self.generate_document_content(content_type, pages, subject)
            print(f' Used standard approach for {pages} pages')
        if file_format == 'Word Document (.docx)':
            file_bytes = self.create_word_document(content)
            suffix = '.docx'
        elif file_format == 'PDF Document (.pdf)':
            file_bytes = self.create_pdf_document(content)
            suffix = '.pdf'
        else:
            file_bytes = self.create_text_document(content)
            suffix = '.txt'
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix, prefix='synthetic_') as tmp_file:
            tmp_file.write(file_bytes)
            temp_path = tmp_file.name
        return (temp_path, f" Generated {pages}-page {content_type} about '{subject}' successfully!")